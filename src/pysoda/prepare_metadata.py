# -*- coding: utf-8 -*-

### Import required python modules

from gevent import monkey; monkey.patch_all()
import platform
import os
from os import listdir, stat, makedirs, mkdir, walk, remove, pardir
from os.path import isdir, isfile, join, splitext, getmtime, basename, normpath, exists, expanduser, split, dirname, getsize, abspath
import pandas as pd
import time
from time import strftime, localtime
import shutil
from shutil import copy2
from configparser import ConfigParser
import numpy as np
from collections import defaultdict
import subprocess
from websocket import create_connection
import socket
import errno
import re
import gevent
from pennsieve import Pennsieve
from pennsieve.log import get_logger
from pennsieve.api.agent import agent_cmd
from pennsieve.api.agent import AgentError, check_port, socket_address
from urllib.request import urlopen
import json
import collections
from threading import Thread
import pathlib

from string import ascii_uppercase
import itertools

from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.styles.colors import Color
from docx import Document

from datetime import datetime, timezone

from Bio import Entrez

userpath = expanduser("~")
metadatapath = join(userpath, 'SODA', 'SODA_metadata')
DEV_TEMPLATE_PATH = join(dirname(__file__), "..", "file_templates")
# once pysoda has been packaged with pyinstaller
# it becomes nested into the pysodadist/api directory
PROD_TEMPLATE_PATH = join(dirname(__file__), "..", "..", "file_templates")
TEMPLATE_PATH = DEV_TEMPLATE_PATH if exists(DEV_TEMPLATE_PATH) else PROD_TEMPLATE_PATH

class InvalidDeliverablesDocument(Exception):
    pass

### Import Milestone document
def import_milestone(filepath):
    doc = Document(filepath)
    try:
        table = doc.tables[0]
    except IndexError:
        raise InvalidDeliverablesDocument("Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: 'Related milestone, aim, or task', 'Description of data', and 'Expected date of completion'.")
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        # headers will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue
        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data

def extract_milestone_info(datalist):
    milestone = defaultdict(list)
    milestone_key1 = "Related milestone, aim, or task"
    milestone_key2 = "Related milestone, aim or task"
    other_keys = ["Description of data", "Expected date of completion"]
    for row in datalist:
        if milestone_key1 in row:
            milestone_key = milestone_key1
        elif milestone_key2 in row:
            milestone_key = milestone_key2
        else:
            raise InvalidDeliverablesDocument("Please select a valid SPARC Deliverables Document! The following headers could not be found in a table of the document you selected: Related milestone, aim, or task, Description of data, and Expected date of completion.")

        key = row[milestone_key]
        if key != "":
            milestone[key].append({key: row[key] for key in other_keys})

    return milestone

### Prepare submission file
def save_submission_file(filepath, json_str):
    source = join(TEMPLATE_PATH, "submission.xlsx")
    destination = filepath
    shutil.copyfile(source, destination)

    # json array to python list
    val_arr = json.loads(json_str)
    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb['Sheet1']
    # date_obj = datetime.strptime(val_arr[2], "%Y-%m")
    # date_new = date_obj.strftime("%m-%Y")
    for column, arr in zip(excel_columns(start_index=2), val_arr):
        ws1[column+"2"] = arr['award']
        ws1[column+"3"] = arr['milestone']
        ws1[column+"4"] = arr['date']

    rename_headers(ws1, len(val_arr), 2)

    wb.save(destination)

def excel_columns(start_index=0):
    """
    NOTE: does not support more than 699 contributors/links
    """
    single_letter = list(ascii_uppercase[start_index:])
    two_letter = [a + b for a,b in itertools.product(ascii_uppercase, ascii_uppercase)]
    return single_letter + two_letter

def rename_headers(workbook, max_len, start_index):
    """
    Rename header columns if values exceed 3. Change Additional Values to Value 4, 5,...
    """

    columns_list = excel_columns(start_index=start_index)
    if max_len >= start_index:

        workbook[columns_list[0] + "1"] = "Value"

        for i, column in zip(range(2, max_len+1), columns_list[1:]):

            workbook[column + "1"] = "Value " + str(i)
            cell = workbook[column + "1"]

            blueFill = PatternFill(start_color='9CC2E5',
                               end_color='9CC2E5',
                               fill_type='solid')

            font = Font(bold=True)
            cell.fill = blueFill
            cell.font = font

    else:

        delete_range = len(columns_list) - max_len
        workbook.delete_cols(4+max_len, delete_range)

### Prepare dataset-description file

def populate_dataset_info(workbook, val_obj):
    ## name, description, type, samples, subjects
    workbook["D5"] = val_obj["name"]
    workbook["D6"] = val_obj["description"]
    workbook["D3"] = val_obj["type"]
    workbook["D29"] = val_obj["number of subjects"]
    workbook["D30"] = val_obj["number of samples"]

    ## keywords
    for i, column in zip(range(len(val_obj["keywords"])), excel_columns(start_index=3)):
        workbook[column + "7"] = val_obj["keywords"][i]

    return val_obj["keywords"]

def populate_study_info(workbook, val_obj):
    workbook["D11"] = val_obj["study purpose"]
    workbook["D12"] = val_obj["study data collection"]
    workbook["D13"] = val_obj["study primary conclusion"]
    workbook["D14"] = val_obj["study organ system"]
    workbook["D15"] = val_obj["study approach"]
    workbook["D16"] = val_obj["study technique"]
    workbook["D17"] = val_obj["study collection title"]

def populate_contributor_info(workbook, val_array):
    ## award info
    for i, column in zip(range(len(val_array["funding"])), excel_columns(start_index=3)):
        workbook[column + "8"] = val_array["funding"][i]

    ### Acknowledgments
    workbook["D9"] = val_array["acknowledgment"]

    ### Contributors
    for contributor, column in zip(val_array['contributors'], excel_columns(start_index=3)):
        workbook[column + "19"] = contributor["conName"]
        workbook[column + "20"] = contributor["conID"]
        workbook[column + "21"] = contributor["conAffliation"]
        # workbook[column + "9"] = contributor["conContact"]
        workbook[column + "22"] = contributor["conRole"]

    return [val_array["funding"], val_array['contributors']]

def populate_related_info(workbook, val_array):
    ## related links including protocols

    for i, column in zip(range(len(val_array)), excel_columns(start_index=3)):
        workbook[column + "24"] = val_array[i]["description"]
        workbook[column + "25"] = val_array[i]["relation"]
        workbook[column + "26"] = val_array[i]["link"]
        workbook[column + "27"] = val_array[i]["type"]

    return len(val_array)


### generate the dataset_description file
import demjson

def save_ds_description_file(bfaccountname, filepath, dataset_str, study_str, con_str, related_info_str):
    source = join(TEMPLATE_PATH, "dataset_description.xlsx")
    destination = filepath
    shutil.copyfile(source, destination)

    # json array to python list
    val_obj_ds = demjson.decode(dataset_str)
    val_obj_study = demjson.decode(study_str)
    val_arr_con = json.loads(con_str)
    val_arr_related_info = json.loads(related_info_str)

    # write to excel file
    wb = load_workbook(destination)
    ws1 = wb['Sheet1']

    keyword_array = populate_dataset_info(ws1, val_obj_ds)

    populate_study_info(ws1, val_obj_study)

    (funding_array, contributor_role_array) = populate_contributor_info(ws1, val_arr_con)

    related_info_len = populate_related_info(ws1, val_arr_related_info)

    # keywords length
    keyword_len = len(keyword_array)

    # contributors length
    no_contributors = len(contributor_role_array)

    # funding = SPARC award + other funding sources
    funding_len = len(funding_array)

    # obtain length for formatting compliance purpose
    max_len = max(keyword_len, funding_len, no_contributors, related_info_len)

    rename_headers(ws1, max_len, 3)

    wb.save(destination)

subjectsTemplateHeaderList = ["subject id", "pool id", "subject experimental group", "age", "sex", "species", "strain", "rrid for strain", "age category", "also in dataset", "member of", "laboratory internal id", "date of birth", "age range (min)", "age range (max)", "body mass", "genotype", "phenotype", "handedness", "reference atlas", "experimental log file path", "experiment date", "disease or disorder", "intervention", "disease model", "protocol title", "protocol url or doi"]
samplesTemplateHeaderList = ["sample id", "subject id", "was derived from", "pool id", "sample experimental group", "sample type", "sample anatomical location", "also in dataset", "member of", "laboratory internal id", "date of derivation", "experimental log file path", "reference atlas", "pathology", "laterality", "cell type", "plane of section", "protocol title", "protocol url or doi"]

def save_subjects_file(filepath, datastructure):

    source = join(TEMPLATE_PATH, "subjects.xlsx")
    destination = filepath
    shutil.copyfile(source, destination)

    wb = load_workbook(destination)
    ws1 = wb['Sheet1']

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:11]
    optionalFields = transposeDatastructure[11:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = subjectsTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(np.concatenate((sortMatrix, refinedOptionalFields)))
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)
    #
    # # 1. delete rows using delete_rows(index, amount=2) -- description and example rows
    # ws1.delete_rows(2, 2)
    # delete all optional columns first (from the template)
    ws1.delete_cols(12, 15)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(start_color='FFD965',
                       end_color='FFD965',
                       fill_type='solid')

    gevent.sleep(0)
    for column, header in zip(excel_columns(start_index=11), refinedDatastructure[0][11:headers_no]):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name='Calibri')

    gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            if refinedDatastructure[i][j]:
                ws1[cell] = refinedDatastructure[i][j]
            else:
                ws1[cell] = ""
            ws1[cell].font = Font(bold=False, size=11, name='Arial')

    wb.save(destination)

def save_samples_file(filepath, datastructure):
    source = join(TEMPLATE_PATH, "samples.xlsx")
    destination = filepath
    shutil.copyfile(source, destination)

    wb = load_workbook(destination)
    ws1 = wb['Sheet1']

    transposeDatastructure = transposeMatrix(datastructure)

    mandatoryFields = transposeDatastructure[:9]
    optionalFields = transposeDatastructure[9:]
    refinedOptionalFields = processMetadataCustomFields(optionalFields)

    templateHeaderList = samplesTemplateHeaderList
    sortMatrix = sortedSubjectsTableData(mandatoryFields, templateHeaderList)

    if refinedOptionalFields:
        refinedDatastructure = transposeMatrix(np.concatenate((sortMatrix, refinedOptionalFields)))
    else:
        refinedDatastructure = transposeMatrix(sortMatrix)

    # # 1. delete rows using delete_rows(index, amount=2) -- description and example rows
    # ws1.delete_rows(2, 2)
    # delete all optional columns first (from the template)
    ws1.delete_cols(10, 9)

    # 2. see if the length of datastructure[0] == length of datastructure. If yes, go ahead. If no, add new columns from headers[n-1] onward.
    headers_no = len(refinedDatastructure[0])
    orangeFill = PatternFill(start_color='FFD965',
                       end_color='FFD965',
                       fill_type='solid')
    gevent.sleep(0)
    for column, header in zip(excel_columns(start_index=9), refinedDatastructure[0][9:headers_no]):
        cell = column + str(1)
        ws1[cell] = header
        ws1[cell].fill = orangeFill
        ws1[cell].font = Font(bold=True, size=12, name='Calibri')

    gevent.sleep(0)
    # 3. populate matrices
    for i, item in enumerate(refinedDatastructure):
        if i == 0:
            continue
        for column, j in zip(excel_columns(start_index=0), range(len(item))):
            # import pdb; pdb.set_trace()
            cell = column + str(i + 1)
            if refinedDatastructure[i][j]:
                ws1[cell] = refinedDatastructure[i][j]
            else:
                ws1[cell] = ""
            ws1[cell].font = Font(bold=False, size=11, name='Arial')

    wb.save(destination)

def column_check(x):
    if 'unnamed' in x.lower():
        return False
    return True

def convert_subjects_samples_file_to_df(type, filepath, ui_fields):
    if type == "subjects":
        templateHeaderList = subjectsTemplateHeaderList
    else:
        templateHeaderList = samplesTemplateHeaderList

    subjects_df = pd.read_excel(filepath, engine='openpyxl', usecols=column_check, header=0)
    subjects_df = subjects_df.dropna(axis = 0, how = 'all')
    subjects_df = subjects_df.applymap(str)
    subjects_df.columns = map(str.lower, subjects_df.columns)
    importedHeaderList = list(subjects_df.columns.values)

    transpose = []
    for header in templateHeaderList:
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""]*len(subjects_df))
        transpose.append(column)

    for header in importedHeaderList:
        if header.lower() in templateHeaderList:
            continue
        column = [header]
        try:
            column.extend(subjects_df[header].values.tolist())
        except KeyError:
            column.extend([""]*len(subjects_df))
        transpose.append(column)

    sortMatrix = sortedSubjectsTableData(transpose, ui_fields)

    return transposeMatrix(sortMatrix)

def sortedSubjectsTableData(matrix, fields):
    sortedMatrix = [];
    customHeaderMatrix = [];

    for field in fields:
        for column in matrix:
            if column[0].lower() == field:
                sortedMatrix.append(column)
                break

    for column in matrix:
        if column[0].lower() not in fields:
            customHeaderMatrix.append(column)

    if len(customHeaderMatrix) > 0:
        npArray = np.concatenate((sortedMatrix, customHeaderMatrix)).tolist()
    else:
        npArray = sortedMatrix
    return npArray

def transposeMatrix(matrix):
    return [[matrix[j][i] for j in range(len(matrix))] for i in range(len(matrix[0]))]

def processMetadataCustomFields(matrix):
    refined_matrix = []
    for column in matrix:
        if any(column[1:]):
            refined_matrix.append(column)

    return refined_matrix

def load_taxonomy_species(animalList):
    animalDict = {}
    for animal in animalList:
        handle = Entrez.esearch(db="Taxonomy", term=animal)
        record = Entrez.read(handle)
        if len(record["IdList"]) > 0:
            id = record['IdList'][0]
            handle = Entrez.efetch(db="Taxonomy", id=id)
            result = Entrez.read(handle)
            animalDict[animal] = {"ScientificName": result[0]['ScientificName'], "OtherNames": result[0]['OtherNames']}

    return animalDict
